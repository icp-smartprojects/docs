#!/usr/bin/env python3
"""
AUREXIS Architecture Document Updater (v2.0)
=============================================
Author: Muriu Mwangi
Date: February 2026

Anchor-based in-place document updater. Updates ONLY sections
between [[SECTION:name]] and [[ENDSECTION:name]] markers.
Preserves all other content. Idempotent — safe to run repeatedly.

Modes:
    python update_doc.py                  # Update existing doc in-place
    python update_doc.py --bootstrap      # Insert anchors into existing doc
    python update_doc.py --new            # Generate fresh doc from scratch
    python update_doc.py --out FILE       # Write to different file
    python update_doc.py --scan           # Scan live services for health/ready

Architecture:
    1. Opens existing AUREXIS_Architecture_Document.docx
    2. Finds [[SECTION:X]] / [[ENDSECTION:X]] anchor pairs
    3. Replaces content between anchors with fresh data
    4. Appends to Changelog (never overwrites history)
    5. Upserts Known Gaps table (match by ID, else append)
    6. Saves in-place (or to --out path)
"""

import os
import sys
import re
import json
import time
import argparse
import subprocess
import hashlib
import io
import math
from pathlib import Path
from datetime import datetime, timezone
from typing import Dict, List, Tuple, Optional, Any
from copy import deepcopy
from dataclasses import dataclass, field

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Paths ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.resolve()
DOC_PATH = BASE_DIR / "AUREXIS_Architecture_Document.docx"
LOGO_PATH = BASE_DIR / "frontend" / "public" / "android-chrome-512x512.png"
COMPOSE_PATH = BASE_DIR / "docker-compose.yml"
TEMP_DIR = BASE_DIR / "temp_doc_images"
TEMP_DIR.mkdir(exist_ok=True)

VERSION = "2.0"
TODAY = datetime.now(timezone.utc).strftime("%Y-%m-%d")

# ─── Colour palette ────────────────────────────────────────────────────
C_PRIMARY    = "#0D1B2A"
C_ACCENT     = "#1B9AAA"
C_SECONDARY  = "#F26419"
C_DATA       = "#1a237e"
C_NERVOUS    = "#1b5e20"
C_PERCEPT    = "#f9a825"
C_COGNITIVE  = "#e65100"
C_DECISION   = "#b71c1c"
C_LEARNING   = "#4a148c"
C_AUDIT      = "#546e7a"
C_INFRA      = "#37474f"
C_FRONTEND   = "#0d47a1"

# ═══════════════════════════════════════════════════════════════════════
# ANCHOR REGISTRY — all managed sections
# ═══════════════════════════════════════════════════════════════════════
MANAGED_SECTIONS = [
    "Service Inventory",
    "Dependency Matrix",
    "Ports and Routing",
    "Event Taxonomy",
    "Enterprise Gate",
    "Known Gaps",
    "Architecture Changelog",
    "Current Status",
    "Next Actions",
    "Improvements Required (62% → 100%)",
    "Runtime Flow Diagrams",
]


# ═══════════════════════════════════════════════════════════════════════
# DATA SOURCES — pull from codebase
# ═══════════════════════════════════════════════════════════════════════

def parse_compose_services() -> Dict[str, Dict]:
    """Parse docker-compose.yml for service definitions."""
    services = {}
    if not COMPOSE_PATH.exists():
        return services

    import yaml
    try:
        with open(COMPOSE_PATH) as f:
            data = yaml.safe_load(f)
        for name, svc in data.get("services", {}).items():
            ports = svc.get("ports", [])
            env = svc.get("environment", {})
            image = svc.get("image", svc.get("build", {}).get("context", ""))
            deps = svc.get("depends_on", [])
            if isinstance(deps, dict):
                deps = list(deps.keys())
            services[name] = {
                "ports": ports,
                "environment": env if isinstance(env, dict) else {},
                "image": str(image),
                "depends_on": deps,
                "healthcheck": svc.get("healthcheck", {}),
            }
    except Exception as e:
        print(f"  ⚠ Could not parse docker-compose.yml: {e}")
    return services


def scan_service_health(services: Dict[str, Dict]) -> Dict[str, Dict]:
    """Attempt to curl /health and /ready for each service."""
    results = {}
    for name, svc in services.items():
        host_port = None
        for p in svc.get("ports", []):
            ps = str(p)
            if ":" in ps:
                host_port = ps.split(":")[0]
                break
        if not host_port:
            results[name] = {"health": "no-port", "ready": "no-port"}
            continue

        health = _curl(f"http://localhost:{host_port}/health")
        ready = _curl(f"http://localhost:{host_port}/ready")
        results[name] = {"health": health, "ready": ready}
    return results


def _curl(url: str, timeout: int = 3) -> str:
    """Quick HTTP GET, returns status or error."""
    try:
        import urllib.request
        req = urllib.request.Request(url, method="GET")
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return f"{resp.status} OK"
    except Exception as e:
        return f"FAIL: {e}"


# ═══════════════════════════════════════════════════════════════════════
# SERVICE INVENTORY — single source of truth
# ═══════════════════════════════════════════════════════════════════════

SERVICE_INVENTORY = [
    {"name": "Event Bus", "port": 52020, "lang": "Go (prod) / Python (dev)", "group": "Nervous System", "docker": "event-bus",
     "inbound": "HTTP POST (publish from all services)", "outbound": "SSE streams to all subscribers", "failure": "ALL streaming stops; services degrade to polling"},
    {"name": "Schema Registry", "port": 52025, "lang": "Go", "group": "Nervous System", "docker": "schema-registry",
     "inbound": "HTTP from Event Bus + API", "outbound": "Validation result; calls Ontology /validate/concepts", "failure": "Unvalidated events published"},
    {"name": "Market Ingestion", "port": 52010, "lang": "Python", "group": "Data Ingestion", "docker": "market-ingestion",
     "inbound": "CSV/Excel files via HTTP upload", "outbound": "batch_committed, market.gap → EB; rows → ClickHouse", "failure": "No new data enters system"},
    {"name": "Candle Constructor", "port": 52023, "lang": "Python", "group": "Data Ingestion", "docker": "candle-constructor",
     "inbound": "ClickHouse reads, EB ticks (live)", "outbound": "LIVE: candles → EB. REPLAY: ClickHouse only", "failure": "No candle stream; perception starves"},
    {"name": "Price Observer", "port": 52002, "lang": "Rust", "group": "Data Ingestion", "docker": "price-observer",
     "inbound": "Raw tick feed", "outbound": "price.observed, tick.received → EB", "failure": "No live ticks"},
    {"name": "Perception", "port": 52012, "lang": "Python", "group": "Perceptual", "docker": "perception",
     "inbound": "candles (SSE), learning.thresholds (SSE)", "outbound": "perception.events → EB", "failure": "No semantic primitives; shape engine blind"},
    {"name": "Shape Engine", "port": 52010, "lang": "Python", "group": "Perceptual", "docker": "shape-engine",
     "inbound": "candles, perception.events, learning.validity (SSE)", "outbound": "shape_created/updated/confirmed/invalidated → EB", "failure": "No shapes; meaning engine starves"},
    {"name": "Meaning Engine", "port": 52003, "lang": "Python", "group": "Cognitive", "docker": "meaning-engine",
     "inbound": "perception.events, shape_* (5 topics), learning.calibration (SSE)", "outbound": "meaning.state, meaning.decisions → EB", "failure": "No beliefs; reasoning blind"},
    {"name": "Reasoning Engine", "port": 52008, "lang": "Python", "group": "Cognitive", "docker": "reasoning-engine",
     "inbound": "meaning.state (SSE)", "outbound": "reasoning.decisions → EB", "failure": "No proofs; core brain decides blind"},
    {"name": "Policy Engine", "port": 52032, "lang": "Python", "group": "Cognitive", "docker": "policy-engine",
     "inbound": "HTTP from CB/SE/LE/SIM; simulation.outcomes, learning.updates (SSE)", "outbound": "policy.allow/deny/evaluated → EB", "failure": "FAIL-CLOSED: all autonomous actions blocked"},
    {"name": "Core Brain", "port": 52040, "lang": "Python", "group": "Decision", "docker": "core-brain",
     "inbound": "13 topics (SSE) + 7 service HTTP calls", "outbound": "brain.decision, decisions → EB", "failure": "No decisions emitted; system observes but doesn't act"},
    {"name": "Simulation", "port": 52043, "lang": "Python", "group": "Decision", "docker": "simulation",
     "inbound": "HTTP trigger from Core Brain", "outbound": "simulation.outcomes, learning.rewards → EB", "failure": "No futures tested; learning starves"},
    {"name": "Learning Engine", "port": 52004, "lang": "Python", "group": "Learning", "docker": "learning-engine",
     "inbound": "11+ topics (SSE)", "outbound": "learning.thresholds/calibration/validity/updates → EB", "failure": "System stops learning; frozen calibration"},
    {"name": "Explanation Engine", "port": 52005, "lang": "Python", "group": "Audit", "docker": "explanation-engine",
     "inbound": "11 topics (SSE)", "outbound": "explanation.ready → EB", "failure": "No explanations; opaque decisions"},
    {"name": "Memory", "port": 52018, "lang": "Python", "group": "Audit", "docker": "memory",
     "inbound": "Events (SSE) + store requests from ME/RE/LE/EE", "outbound": "HTTP recall responses to ME/RE/LE/EE", "failure": "No experience recall; cognition loop breaks"},
    {"name": "Knowledge Graph", "port": 52015, "lang": "Python", "group": "Audit", "docker": "knowledge-graph",
     "inbound": "SSE: 5 EB topics + HTTP CRUD", "outbound": "kg.node_created/edge_created → EB", "failure": "No semantic graph; ontology validation degrades"},
    {"name": "Language Intelligence", "port": 52006, "lang": "Python", "group": "Audit", "docker": "language-intelligence",
     "inbound": "HTTP NL queries + explanation.ready (SSE)", "outbound": "HTTP NL responses", "failure": "No natural language interface"},
    {"name": "Ontology", "port": 52100, "lang": "Python", "group": "Audit", "docker": "ontology",
     "inbound": "HTTP CRUD + validation requests from Perc/SE/ME/PE/SR/KG", "outbound": "HTTP validation responses", "failure": "Services use unchecked concept names"},
    {"name": "Gateway", "port": 52031, "lang": "Go (prod) / Python (dev)", "group": "Infrastructure", "docker": "gateway",
     "inbound": "Frontend HTTP + WS", "outbound": "Proxied to 17+ downstream services", "failure": "Frontend cannot reach backend (control plane down)"},
    {"name": "Topology Hub", "port": 8080, "lang": "Go", "group": "Infrastructure", "docker": "topology-hub",
     "inbound": "Service heartbeats", "outbound": "Topology graph via WebSocket", "failure": "No service visibility"},
    {"name": "Topology Hub Express", "port": 3000, "lang": "Node.js", "group": "Infrastructure", "docker": "topology-hub-express",
     "inbound": "Frontend Socket.io", "outbound": "Proxied to Topology Hub", "failure": "Frontend loses topology view"},
    {"name": "Schema Registry", "port": 52025, "lang": "Go", "group": "Infrastructure", "docker": "schema-registry",
     "inbound": "Schema registration + validation", "outbound": "Validation verdicts; Ontology concept checks", "failure": "Unvalidated schemas accepted"},
    {"name": "Auth System", "port": 5000, "lang": "Node.js", "group": "Infrastructure", "docker": "auth-system",
     "inbound": "Login/register/2FA HTTP", "outbound": "JWT tokens", "failure": "No authentication; all requests rejected"},
    {"name": "Security Core", "port": 55000, "lang": "Node.js", "group": "Infrastructure", "docker": "security",
     "inbound": "Crypto/audit/ABAC HTTP", "outbound": "Crypto results, audit log", "failure": "No encryption, no audit trail"},
    {"name": "Frontend", "port": 80, "lang": "React 18 + TypeScript", "group": "Infrastructure", "docker": "frontend",
     "inbound": "HTTP from Gateway + SSE/WS from EB", "outbound": "User HTTP requests → Gateway", "failure": "No UI"},
]

# Deduplicate by (name, port) — Schema Registry appears twice in source data
_seen = set()
_deduped = []
for s in SERVICE_INVENTORY:
    key = (s["name"], s["port"])
    if key not in _seen:
        _seen.add(key)
        _deduped.append(s)
SERVICE_INVENTORY = _deduped


# ═══════════════════════════════════════════════════════════════════════
# CANONICAL EVENT TAXONOMY
# ═══════════════════════════════════════════════════════════════════════

EVENT_TAXONOMY = [
    # [Layer, Topic, Publisher, Subscribers, Schema Version]
    ["EXTERNAL", "price.observed", "Price Observer", "Perception, Memory", "v1"],
    ["EXTERNAL", "candles (LIVE only)", "Candle Constructor", "Perception, Shape Engine", "v1"],
    ["EXTERNAL", "tick.received", "Price Observer", "Perception", "v1"],
    ["EXTERNAL", "market.ingestion.batch_committed", "Market Ingestion", "Candle Constructor", "v1"],
    ["EXTERNAL", "market.gap", "Market Ingestion", "Memory", "v1"],
    ["PERCEPTUAL", "perception.events", "Perception", "Shape Engine, Meaning Engine", "v1"],
    ["COGNITIVE", "shape_created", "Shape Engine", "Meaning, Core Brain, Learning, KG", "v1"],
    ["COGNITIVE", "shape_updated", "Shape Engine", "Meaning, Learning", "v1"],
    ["COGNITIVE", "shape_confirmed", "Shape Engine", "Meaning, Core Brain, Learning, KG", "v1"],
    ["COGNITIVE", "shape_invalidated", "Shape Engine", "Meaning, Core Brain, Learning", "v1"],
    ["COGNITIVE", "shape_touched", "Shape Engine", "Meaning, Learning", "v1"],
    ["COGNITIVE", "meaning.state", "Meaning Engine", "Reasoning, Core Brain, Learning, KG", "v1"],
    ["COGNITIVE", "meaning.decisions", "Meaning Engine", "Core Brain", "v1"],
    ["COGNITIVE", "reasoning.decisions", "Reasoning Engine", "Core Brain, Learning, KG", "v1"],
    ["COGNITIVE", "ontology.updated", "Ontology", "KG", "v1"],
    ["DECISION", "policy.allow", "Policy Engine", "Core Brain", "v1"],
    ["DECISION", "policy.deny", "Policy Engine", "Core Brain", "v1"],
    ["DECISION", "policy.evaluated", "Policy Engine", "Explanation Engine", "v1"],
    ["DECISION", "brain.decision", "Core Brain", "Explanation, Learning, Policy", "v1"],
    ["DECISION", "decisions", "Core Brain", "Explanation, Language Intel", "v1"],
    ["OUTCOME", "simulation.outcomes", "Simulation", "Learning, Policy, Core Brain", "v1"],
    ["LEARNING", "learning.thresholds", "Learning Engine", "Perception", "v1"],
    ["LEARNING", "learning.calibration", "Learning Engine", "Meaning Engine", "v1"],
    ["LEARNING", "learning.validity", "Learning Engine", "Shape Engine", "v1"],
    ["LEARNING", "learning.updates", "Learning Engine", "Core Brain, Policy", "v1"],
    ["SYSTEM", "explanation.ready", "Explanation Engine", "Language Intelligence", "v1"],
]


# ═══════════════════════════════════════════════════════════════════════
# KNOWN GAPS TABLE (ID-based for upsert)
# ═══════════════════════════════════════════════════════════════════════

KNOWN_GAPS = [
    {"id": "GAP-001", "service": "Event Bus / Gateway", "severity": "MEDIUM",
     "status": "OPEN", "summary": "Dual Go+Python implementations create confusion. Dockerfile builds Go, but Python stubs remain.",
     "fix_plan": "Archive Python stubs to /archive/, remove from COPY in Dockerfile"},
    {"id": "GAP-002", "service": "ALL", "severity": "HIGH",
     "status": "OPEN", "summary": "K8s manifests and Helm charts are all empty files. No Kubernetes deployment possible.",
     "fix_plan": "Generate from docker-compose using kompose + manual tuning"},
    {"id": "GAP-003", "service": "Event Bus", "severity": "MEDIUM",
     "status": "OPEN", "summary": "Protobuf serialization returns 'not implemented'. JSON-only.",
     "fix_plan": "Implement protobuf codec or remove proto references"},
    {"id": "GAP-004", "service": "Candle Constructor", "severity": "CRITICAL",
     "status": "FIXED", "summary": "SQL injection risk in aggregate_ticks_to_candles() — f-string interpolation of symbol/tf.",
     "fix_plan": "FIXED: Added allowlist sanitisation regex + quote stripping"},
    {"id": "GAP-005", "service": "Ontology", "severity": "HIGH",
     "status": "FIXED", "summary": "eval() usage in constraint_engine.py — sandbox escapable.",
     "fix_plan": "FIXED: Replaced with AST node whitelist + compile() on validated tree"},
    {"id": "GAP-006", "service": "Learning Engine", "severity": "HIGH",
     "status": "FIXED", "summary": "CORS allow_origins=['*'] with allow_credentials=True — spec violation.",
     "fix_plan": "FIXED: Fallback changed to explicit localhost origins"},
    {"id": "GAP-007", "service": "Gateway", "severity": "CRITICAL",
     "status": "FIXED", "summary": "engine_controller.py used localhost for all 20 service health URLs — fails inside Docker.",
     "fix_plan": "FIXED: All URLs now use Docker service DNS names"},
    {"id": "GAP-008", "service": "Gateway", "severity": "CRITICAL",
     "status": "FIXED", "summary": "gateway.yaml had wrong ports (4222, 5001-5009) + localhost for all services.",
     "fix_plan": "FIXED: Corrected to Docker DNS names + actual ports"},
    {"id": "GAP-009", "service": "Gateway", "severity": "HIGH",
     "status": "FIXED", "summary": "Go source files (eventbus.go, policy.go, api.go) had localhost fallback defaults.",
     "fix_plan": "FIXED: Defaults changed to Docker service DNS"},
    {"id": "GAP-010", "service": "Learning Engine", "severity": "CRITICAL",
     "status": "FIXED", "summary": "MEMORY_URL pointed to port 52011 (shape-engine gRPC) instead of 52018 (memory).",
     "fix_plan": "FIXED: docker-compose.yml corrected to http://memory:52018"},
    {"id": "GAP-011", "service": "Market Ingestion", "severity": "MEDIUM",
     "status": "OPEN", "summary": "API/Broker/WebSocket live data connectors all stubbed. CSV-only ingestion.",
     "fix_plan": "Implement at least one live connector (WebSocket broker)"},
    {"id": "GAP-012", "service": "Policy Engine", "severity": "LOW",
     "status": "OPEN", "summary": "Harm detector always returns True, ethical validator always True.",
     "fix_plan": "Implement rule-based harm/ethical checks or defer to config"},
    {"id": "GAP-013", "service": "Simulation", "severity": "LOW",
     "status": "OPEN", "summary": "_transform_timeframe() returns empty list.",
     "fix_plan": "Implement TF transformation or remove dead code"},
    {"id": "GAP-014", "service": "Frontend", "severity": "MEDIUM",
     "status": "OPEN", "summary": "Sidebar panels for reasoning, policy, core brain not wired to backend endpoints.",
     "fix_plan": "Wire API calls to existing service endpoints"},
    {"id": "GAP-015", "service": "Frontend", "severity": "MEDIUM",
     "status": "OPEN", "summary": "Training UI missing accept/reject/adjust feedback loop for auto-shapes.",
     "fix_plan": "Add shape feedback UI → POST to learning-engine"},
    {"id": "GAP-016", "service": "Meaning Engine", "severity": "LOW",
     "status": "OPEN", "summary": "Confluence score and contradiction flags not exposed on UI.",
     "fix_plan": "Add sidebar endpoint + frontend panel"},
]

IMPROVEMENTS = [
    {"priority": 0, "id": "IMP-P0-01", "title": "Dual implementations (Go vs Python)",
     "desc": "event-bus/ and gateway/ contain both Go production code and Python dev stubs. Pick one runtime, archive the other."},
    {"priority": 0, "id": "IMP-P0-02", "title": "Port mismatches enforced",
     "desc": "FIXED: engine_controller.py ports corrected, MEMORY_URL fixed, gateway.yaml ports corrected."},
    {"priority": 0, "id": "IMP-P0-03", "title": "ClickHouse query injection sanitised",
     "desc": "FIXED: aggregate_ticks_to_candles() now uses allowlist regex validation before SQL interpolation."},
    {"priority": 0, "id": "IMP-P0-04", "title": "eval() removed from constraint engine",
     "desc": "FIXED: Replaced with AST whitelist validator + compile() — no raw eval() in codebase."},
    {"priority": 0, "id": "IMP-P0-05", "title": "CORS wildcard with credentials fixed",
     "desc": "FIXED: learning-engine fallback origins changed from ['*'] to explicit localhost list."},
    {"priority": 0, "id": "IMP-P0-06", "title": "No-localhost rule enforced in Docker",
     "desc": "FIXED: All gateway refs (engine_controller.py, Go sources, .env, gateway.yaml) now use Docker service DNS."},
    {"priority": 0, "id": "IMP-P0-07", "title": "Missing /ready endpoints",
     "desc": "VERIFIED: All 24 services already have /ready endpoints."},
    {"priority": 1, "id": "IMP-P1-01", "title": "Schema registry enforcement",
     "desc": "Versioning + backward compatibility gates needed. Schema lifecycle (draft→active→deprecated) exists but not enforced on publish."},
    {"priority": 1, "id": "IMP-P1-02", "title": "Dead-letter queue / retry discipline",
     "desc": "Event Bus Python stub has DLQ; Go production binary does not. Needs consistent implementation."},
    {"priority": 1, "id": "IMP-P1-03", "title": "Deterministic replay for simulation",
     "desc": "SHA-256 checksums exist but full replay determinism not verified end-to-end."},
    {"priority": 1, "id": "IMP-P1-04", "title": "Observability everywhere",
     "desc": "Prometheus /metrics exists on most services but format varies. Tracing (OpenTelemetry) not wired. Structured logging inconsistent."},
    {"priority": 1, "id": "IMP-P1-05", "title": "Gate tests per service",
     "desc": "Minimum: health, ready, OpenAPI, basic load. Currently no automated test suite."},
    {"priority": 1, "id": "IMP-P1-06", "title": "Remove committed artifacts",
     "desc": "DB files, compiled binaries, audit logs, *.bak files committed to repo. Need .gitignore cleanup."},
    {"priority": 2, "id": "IMP-P2-01", "title": "Frontend sidebar wiring",
     "desc": "Reasoning, policy, core brain sidebar panels exist but API calls not connected."},
    {"priority": 2, "id": "IMP-P2-02", "title": "Training feedback UI",
     "desc": "Accept/reject/adjust loop for auto-shapes not implemented in frontend."},
    {"priority": 2, "id": "IMP-P2-03", "title": "Confluence + contradiction display",
     "desc": "Confluence score and contradiction flags from Meaning Engine not shown on UI."},
    {"priority": 2, "id": "IMP-P2-04", "title": "Skeleton mode",
     "desc": "Time-space geometry view (wireframe candles, shape anchors, swing vectors) not implemented in chart layer."},
]


# ═══════════════════════════════════════════════════════════════════════
# DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, insert_before=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "0D1B2A")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7)
            bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
            set_cell_shading(cell, bg)
    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(13, 27, 42)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(27, 154, 170)
    return p


# ═══════════════════════════════════════════════════════════════════════
# ANCHOR SYSTEM — find/replace between markers
# ═══════════════════════════════════════════════════════════════════════

def find_anchor_range(doc, section_name: str) -> Tuple[Optional[int], Optional[int]]:
    """
    Find paragraph indexes for [[SECTION:name]] and [[ENDSECTION:name]].
    Returns (start_idx, end_idx) or (None, None) if not found.
    """
    start_marker = f"[[SECTION:{section_name}]]"
    end_marker = f"[[ENDSECTION:{section_name}]]"
    start_idx = None
    end_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == start_marker:
            start_idx = i
        elif text == end_marker and start_idx is not None:
            end_idx = i
            break

    return start_idx, end_idx


def clear_between_anchors(doc, start_idx: int, end_idx: int):
    """Remove all paragraphs between start and end anchors (exclusive)."""
    # We need to remove paragraphs from end_idx-1 down to start_idx+1
    body = doc.element.body
    paras_to_remove = []
    for i in range(start_idx + 1, end_idx):
        paras_to_remove.append(doc.paragraphs[i]._element)

    for elem in paras_to_remove:
        body.remove(elem)


def insert_anchor_pair(doc, section_name: str, after_paragraph_idx: int = None):
    """Insert [[SECTION:X]] and [[ENDSECTION:X]] markers."""
    if after_paragraph_idx is not None:
        # Insert after specific paragraph
        ref_elem = doc.paragraphs[after_paragraph_idx]._element
        start_p = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:t>[[SECTION:{section_name}]]</w:t></w:r></w:p>')
        end_p = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:t>[[ENDSECTION:{section_name}]]</w:t></w:r></w:p>')
        ref_elem.addnext(end_p)
        ref_elem.addnext(start_p)
    else:
        # Append to end
        p1 = doc.add_paragraph(f"[[SECTION:{section_name}]]")
        p1.runs[0].font.size = Pt(1)
        p1.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p2 = doc.add_paragraph(f"[[ENDSECTION:{section_name}]]")
        p2.runs[0].font.size = Pt(1)
        p2.runs[0].font.color.rgb = RGBColor(255, 255, 255)


# ═══════════════════════════════════════════════════════════════════════
# DIAGRAM HELPERS (reused from v1)
# ═══════════════════════════════════════════════════════════════════════

def _save_fig(fig, name):
    p = TEMP_DIR / f"{name}.png"
    fig.savefig(str(p), dpi=200, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return str(p)

def _draw_box(ax, x, y, w, h, text, color="#1B9AAA", text_color="white", fontsize=7, alpha=1.0):
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                         facecolor=color, edgecolor="white", linewidth=0.8, alpha=alpha)
    ax.add_patch(box)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        fs = fontsize if i > 0 else fontsize + 1
        fw = "bold" if i == 0 else "normal"
        ax.text(x + w/2, y + h - 0.25 - i*0.35, line,
                ha="center", va="center", fontsize=fs, color=text_color, fontweight=fw)

def _draw_arrow(ax, x1, y1, x2, y2, color="#AAAAAA", style="-|>", lw=1.0, ls="-"):
    arrow = FancyArrowPatch((x1, y1), (x2, y2),
                            arrowstyle=style, color=color,
                            linewidth=lw, linestyle=ls,
                            mutation_scale=10,
                            connectionstyle="arc3,rad=0.0")
    ax.add_patch(arrow)

def _draw_label(ax, x1, y1, x2, y2, text, color="#CCCCCC", fontsize=5):
    mx, my = (x1+x2)/2, (y1+y2)/2
    ax.text(mx, my + 0.12, text, ha="center", va="bottom",
            fontsize=fontsize, color=color, style="italic")


def make_data_plane_diagram():
    """Create runtime data plane diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    fig.patch.set_facecolor("#0D1B2A")
    ax.set_facecolor("#0D1B2A")
    ax.set_xlim(-1, 15)
    ax.set_ylim(-0.5, 11)
    ax.axis("off")

    ax.text(7, 10.5, "DATA PLANE — Runtime Service Flow", ha="center",
            fontsize=14, color="white", fontweight="bold")

    # Market Ingestion → ClickHouse
    _draw_box(ax, 0, 8.5, 2.5, 1.2, "Market Ingestion\n:52010\nCSV → ClickHouse", color=C_DATA, fontsize=6)
    _draw_box(ax, 3.5, 8.5, 2.5, 1.2, "ClickHouse\nCandle Truth\nSource", color="#0D47A1", fontsize=6)
    _draw_arrow(ax, 2.5, 9.1, 3.5, 9.1, color="#90CAF9")
    _draw_label(ax, 2.5, 9.1, 3.5, 9.1, "write OHLCV")

    # CC + PO → Event Bus
    _draw_box(ax, 7, 8.5, 3, 1.2, "Candle Constructor\n:52023\nLIVE → EB", color=C_DATA, fontsize=6)
    _draw_box(ax, 11, 8.5, 2.5, 1.2, "Price Observer\n:52002 · Rust", color="#1A237E", fontsize=6)
    _draw_arrow(ax, 6, 9.1, 7, 9.1, color="#90CAF9")

    # Event Bus spine
    eb_y = 6.5
    rect = FancyBboxPatch((1, eb_y), 12, 1.0, boxstyle="round,pad=0.15",
                          facecolor=C_NERVOUS, edgecolor="#4CAF50", linewidth=2, alpha=0.5)
    ax.add_patch(rect)
    ax.text(7, eb_y + 0.5, "EVENT BUS  —  DATA PLANE  |  :52020  |  SSE + HTTP",
            ha="center", va="center", fontsize=9, color="#A5D6A7", fontweight="bold")

    _draw_arrow(ax, 8.5, 8.5, 7, eb_y + 1.0, color="#A5D6A7")
    _draw_label(ax, 8.5, 8.5, 7, eb_y+1.0, "candles.stream")
    _draw_arrow(ax, 12.25, 8.5, 12, eb_y + 1.0, color="#A5D6A7")

    # Perceptual
    _draw_box(ax, 1, 4.5, 3, 1.2, "Perception\n:52012\nperception.events →", color="#F57F17", fontsize=6)
    _draw_box(ax, 5, 4.5, 3.5, 1.2, "Shape Engine\n:52010\nshape.* events →", color=C_COGNITIVE, fontsize=6)
    _draw_arrow(ax, 2.5, eb_y, 2.5, 5.7, color="#FFF176")
    _draw_arrow(ax, 6.75, eb_y, 6.75, 5.7, color="#FFCC80")

    # Cognitive
    _draw_box(ax, 0, 2.5, 3, 1.2, "Meaning Engine\n:52003\nmeaning.state →", color="#BF360C", fontsize=6)
    _draw_box(ax, 4, 2.5, 3, 1.2, "Reasoning Engine\n:52008\nreasoning.proofs →", color="#D84315", fontsize=6)
    _draw_arrow(ax, 1.5, 4.5, 1.5, 3.7, color="#FFCC80", ls="--")
    _draw_arrow(ax, 5.5, 4.5, 5.5, 3.7, color="#FFCC80", ls="--")

    # Decision
    _draw_box(ax, 8.5, 2.5, 3, 1.2, "Core Brain\n:52040\nbrain.decision →", color=C_DECISION, fontsize=6)
    _draw_box(ax, 12, 2.5, 2.5, 1.2, "Policy Engine\n:52032\nFail-Closed", color="#E64A19", fontsize=6)
    _draw_arrow(ax, 7, 3.1, 8.5, 3.1, color="white")
    _draw_arrow(ax, 11.5, 3.1, 12, 3.1, color="white")

    # Simulation + Learning
    _draw_box(ax, 8.5, 0.5, 3, 1.2, "Simulation\n:52043\nsim.outcomes →", color="#880E4F", fontsize=6)
    _draw_box(ax, 1, 0.5, 4, 1.2, "Learning Engine\n:52004\nthresholds/calibration/validity →", color=C_LEARNING, fontsize=6)
    _draw_arrow(ax, 10, 2.5, 10, 1.7, color="#EF9A9A")
    # Feedback arrow from learning up
    _draw_arrow(ax, 3, 1.7, 2.5, 4.5, color="#CE93D8", ls="--", lw=1.5)

    # Memory + EE
    _draw_box(ax, 12, 0.5, 2.5, 1.2, "Memory\n:52018\n→ ME/RE/LE/EE", color=C_AUDIT, fontsize=6)

    return _save_fig(fig, "data_plane")


def make_control_plane_diagram():
    """Create runtime control plane diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(12, 6))
    fig.patch.set_facecolor("#0D1B2A")
    ax.set_facecolor("#0D1B2A")
    ax.set_xlim(-0.5, 13)
    ax.set_ylim(-0.5, 6)
    ax.axis("off")

    ax.text(6.5, 5.5, "CONTROL PLANE — Frontend → Gateway → Services", ha="center",
            fontsize=13, color="white", fontweight="bold")

    # Frontend
    _draw_box(ax, 0, 2, 2.5, 1.5, "Frontend\n:80 React SPA\nSSE + WS Live", color=C_FRONTEND, fontsize=7)

    # Gateway
    _draw_box(ax, 4, 1.5, 3, 2.5, "Gateway\n:52031 · Go\nCONTROL PLANE\nReverse Proxy\nJWT · CORS · 17 routes", color=C_NERVOUS, fontsize=6)
    _draw_arrow(ax, 2.5, 2.75, 4, 2.75, color="#64B5F6")
    _draw_label(ax, 2.5, 2.75, 4, 2.75, "HTTP /api/*")

    # Auth
    _draw_box(ax, 8.5, 4, 3, 1.0, "Auth System :5000", color="#263238", fontsize=7)
    _draw_arrow(ax, 7, 3.5, 8.5, 4.5, color="#90A4AE")
    _draw_label(ax, 7, 3.5, 8.5, 4.5, "JWT verify")

    # Downstream services
    services_right = [
        "Candle Constructor :52023",
        "Shape Engine :52010",
        "Meaning Engine :52003",
        "Reasoning Engine :52008",
        "Simulation :52043",
        "Learning Engine :52004",
        "Explanation Engine :52005",
        "Memory :52018",
        "Ontology :52100",
    ]
    for i, svc in enumerate(services_right):
        y = 0.2 + i * 0.48
        _draw_box(ax, 8.5, y, 4, 0.4, svc, color="#263238", fontsize=5)
        _draw_arrow(ax, 7, 2.75, 8.5, y + 0.2, color="#78909C", lw=0.5)

    # Topology
    _draw_box(ax, 0, 0, 2.5, 1.2, "Topology Express\n:3000 · Socket.io", color="#263238", fontsize=6)
    _draw_arrow(ax, 1.25, 2.0, 1.25, 1.2, color="#CE93D8", ls="--")
    _draw_label(ax, 1.25, 2.0, 1.25, 1.2, "socket")

    return _save_fig(fig, "control_plane")


# ═══════════════════════════════════════════════════════════════════════
# SECTION GENERATORS — produce content for each anchor
# ═══════════════════════════════════════════════════════════════════════

def generate_service_inventory(doc):
    """Generate Service Inventory table."""
    add_heading_styled(doc, "Service Inventory (24 Services)", 2)
    rows = []
    for s in SERVICE_INVENTORY:
        rows.append([s["name"], str(s["port"]), s["lang"], s["group"], s["docker"]])
    add_styled_table(doc, ["Service", "Port", "Language", "Group", "Docker Name"], rows)
    doc.add_paragraph()


def generate_dependency_matrix(doc):
    """Generate Runtime Dependency Matrix."""
    add_heading_styled(doc, "Runtime Dependency Matrix", 2)
    add_body(doc, "What each service consumes (inbound), publishes (outbound), and what happens when it fails.")
    rows = []
    for s in SERVICE_INVENTORY:
        rows.append([s["name"], s.get("inbound", "—"), s.get("outbound", "—"), s.get("failure", "—")])
    add_styled_table(doc, ["Service", "Inbound (Consumes)", "Outbound (Publishes)", "Failure Mode"], rows)
    doc.add_paragraph()


def generate_ports_routing(doc, compose_services: Dict):
    """Generate Ports & Routing table from docker-compose."""
    add_heading_styled(doc, "Ports & Routing (from docker-compose.yml)", 2)
    rows = []
    for name, svc in sorted(compose_services.items()):
        ports = ", ".join(str(p) for p in svc.get("ports", []))
        deps = ", ".join(svc.get("depends_on", []))
        hc = svc.get("healthcheck", {})
        hc_cmd = " ".join(hc.get("test", ["—"])[1:]) if isinstance(hc.get("test"), list) else "—"
        rows.append([name, ports or "—", deps or "—", hc_cmd[:60]])
    add_styled_table(doc, ["Service", "Ports (host:container)", "depends_on", "Healthcheck"], rows)
    doc.add_paragraph()


def generate_event_taxonomy(doc):
    """Generate Event Taxonomy table."""
    add_heading_styled(doc, "Canonical Event Taxonomy (26 Topics)", 2)
    add_body(doc, "Single source of truth for all Event Bus topics. Schema version column tracks contract evolution.")
    add_styled_table(doc, ["Layer", "Topic", "Publisher", "Subscribers", "Schema"], EVENT_TAXONOMY)
    doc.add_paragraph()


def generate_enterprise_gate(doc, scan_results: Optional[Dict]):
    """Generate Enterprise Gate test results."""
    add_heading_styled(doc, "Enterprise Gate Tests", 2)
    if scan_results:
        rows = []
        for name, res in sorted(scan_results.items()):
            health = res.get("health", "—")
            ready = res.get("ready", "—")
            health_ok = "✅" if "200" in str(health) else "❌"
            ready_ok = "✅" if "200" in str(ready) else "❌"
            rows.append([name, f"{health_ok} {health}", f"{ready_ok} {ready}"])
        add_styled_table(doc, ["Service", "/health", "/ready"], rows)
    else:
        add_body(doc, "Runtime checks skipped (services offline or --scan not used).")
    doc.add_paragraph()


def generate_known_gaps(doc):
    """Generate Known Gaps table (ID-based for upsert)."""
    add_heading_styled(doc, "Known Gaps & Defects", 2)
    rows = []
    for g in KNOWN_GAPS:
        rows.append([g["id"], g["service"], g["severity"], g["status"], g["summary"][:80], g["fix_plan"][:60]])
    add_styled_table(doc, ["ID", "Service", "Severity", "Status", "Summary", "Fix Plan"], rows)
    doc.add_paragraph()


def generate_changelog(doc, existing_entries: List[str] = None):
    """Generate Architecture Changelog (append-only)."""
    add_heading_styled(doc, "Architecture Changelog", 2)

    # Preserve existing entries
    if existing_entries:
        for entry in existing_entries:
            add_body(doc, entry)

    # Add today's entry if not already present
    today_entry = f"[{TODAY}] v{VERSION} — Priority 0 fixes applied: SQL injection sanitised, eval() replaced with AST validator, CORS wildcard removed, gateway localhost→Docker DNS (20 services), MEMORY_URL port corrected (52011→52018), engine_controller.py ports fixed. Anchor-based document updater introduced."
    if not existing_entries or not any(TODAY in e for e in existing_entries):
        add_body(doc, today_entry)

    doc.add_paragraph()


def generate_current_status(doc):
    """Generate Current Status percentage."""
    add_heading_styled(doc, "Current System Status", 2)

    scores = {
        "Service coverage": 80,
        "Runtime wiring correctness": 72,  # was 55, fixed ports/DNS/MEMORY_URL
        "Contract discipline (topics/schemas)": 50,  # was 45, added taxonomy
        "Observability (metrics/tracing/readiness)": 58,  # was 55, /ready verified
        "Frontend integration completeness": 60,
        "Security posture": 70,  # was ~50, fixed eval/CORS/injection
    }

    overall = sum(scores.values()) / len(scores)

    rows = [[k, f"{v}/100"] for k, v in scores.items()]
    rows.append(["OVERALL", f"{overall:.0f}/100"])
    add_styled_table(doc, ["Dimension", "Score"], rows)
    doc.add_paragraph()
    add_body(doc, f"Overall score: {overall:.0f}% (up from 62% before P0 fixes).")
    doc.add_paragraph()


def generate_next_actions(doc):
    """Generate Next Actions / TODOs."""
    add_heading_styled(doc, "Next Actions", 2)
    actions = [
        "1. Archive Python stubs from event-bus/ and gateway/ (dual impl cleanup)",
        "2. Implement schema versioning enforcement gates on Event Bus publish",
        "3. Add dead-letter queue to Go Event Bus binary",
        "4. Wire frontend sidebar panels to reasoning/policy/core-brain endpoints",
        "5. Add training feedback UI (accept/reject/adjust auto-shapes)",
        "6. Implement at least one live data connector (WebSocket broker)",
        "7. Add OpenTelemetry tracing across all Python services",
        "8. Generate K8s manifests from docker-compose (kompose + manual tuning)",
        "9. Set up CI/CD gate tests: health, ready, OpenAPI, basic load per service",
        "10. Implement skeleton mode in frontend chart layer",
    ]
    for a in actions:
        doc.add_paragraph(a, style="List Bullet")
    doc.add_paragraph()


def generate_improvements(doc):
    """Generate Improvements Required (62%→100%) section."""
    add_heading_styled(doc, "Improvements Required to Reach 100%", 2)

    for priority in [0, 1, 2]:
        label = {0: "Priority 0 — Must Fix (Security & Runtime)", 1: "Priority 1 — Enterprise Reliability", 2: "Priority 2 — Product Completeness"}[priority]
        add_heading_styled(doc, label, 3)
        items = [i for i in IMPROVEMENTS if i["priority"] == priority]
        rows = [[i["id"], i["title"], "FIXED" if "FIXED" in i["desc"] else "OPEN", i["desc"][:80]] for i in items]
        add_styled_table(doc, ["ID", "Title", "Status", "Description"], rows)
        doc.add_paragraph()


def generate_runtime_flow(doc):
    """Generate Data Plane + Control Plane runtime flow diagrams."""
    add_heading_styled(doc, "Runtime Flow Diagrams", 2)

    add_heading_styled(doc, "A) Data Plane (Speed + Cognition)", 3)
    add_body(doc, "Services communicate through the Event Bus (SSE/HTTP). This is the east–west nervous system for real-time streaming events.")
    img_data = make_data_plane_diagram()
    doc.add_picture(img_data, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_styled(doc, "B) Control Plane (UI + Access)", 3)
    add_body(doc, "Frontend reaches services through the Gateway reverse proxy. This is the north–south API routing layer with auth, CORS, and rate limiting.")
    img_ctrl = make_control_plane_diagram()
    doc.add_picture(img_ctrl, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_styled(doc, "The Rule", 3)
    add_body(doc, "If it's a command or query → Gateway/HTTP (control plane). If it's a stream of facts/events → Event Bus (data plane).")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# BOOTSTRAP MODE — insert anchors into existing doc
# ═══════════════════════════════════════════════════════════════════════

def bootstrap_anchors(doc_path: Path) -> Path:
    """Insert anchor markers into existing doc."""
    print(f"Bootstrapping anchors into {doc_path}...")
    doc = Document(str(doc_path))

    # Find the last paragraph
    existing_anchors = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        m = re.match(r'\[\[SECTION:(.+?)\]\]', text)
        if m:
            existing_anchors.add(m.group(1))

    # Add page break before new sections
    doc.add_page_break()
    add_heading_styled(doc, "Part XI — Enterprise Qualification Sections", 1)
    add_body(doc, "The following sections are managed by the anchor-based document updater (update_doc.py). They are automatically refreshed on each run.")

    for section_name in MANAGED_SECTIONS:
        if section_name in existing_anchors:
            print(f"  ⏭ Anchor already exists: {section_name}")
            continue
        print(f"  ✅ Inserting anchor: [[SECTION:{section_name}]]")
        doc.add_paragraph()
        add_heading_styled(doc, section_name.replace(" (62% → 100%)", ""), 2)
        p1 = doc.add_paragraph(f"[[SECTION:{section_name}]]")
        for run in p1.runs:
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(180, 180, 180)
        p2 = doc.add_paragraph(f"[[ENDSECTION:{section_name}]]")
        for run in p2.runs:
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(180, 180, 180)

    doc.save(str(doc_path))
    print(f"Anchors bootstrapped. Saved to {doc_path}")
    return doc_path


# ═══════════════════════════════════════════════════════════════════════
# UPDATE MODE — replace content between anchors
# ═══════════════════════════════════════════════════════════════════════

def update_section(doc, section_name: str, generator_fn, *args):
    """
    Find anchors, clear content between them, generate new content.
    Returns True if section was updated, False if anchor not found.
    """
    start_idx, end_idx = find_anchor_range(doc, section_name)
    if start_idx is None or end_idx is None:
        print(f"  ⚠ Anchor not found: [[SECTION:{section_name}]] — skipping")
        return False

    # Clear existing content
    clear_between_anchors(doc, start_idx, end_idx)

    # The generator appends to doc.paragraphs, but we need to insert
    # BETWEEN the anchors. Since we cleared everything between them,
    # the end anchor is now at start_idx + 1.
    # We'll use a temporary document approach:
    temp_doc = Document()
    generator_fn(temp_doc, *args)

    # Insert temp doc content before the end anchor
    end_anchor_elem = doc.paragraphs[start_idx + 1]._element
    body = doc.element.body

    # Collect all elements from temp doc
    temp_elements = list(temp_doc.element.body)
    for elem in reversed(temp_elements):
        # Deep copy the element
        new_elem = deepcopy(elem)
        end_anchor_elem.addprevious(new_elem)

    return True


def update_document(doc_path: Path, out_path: Path = None, do_scan: bool = False):
    """Main update routine — open doc, update anchored sections, save."""
    if not doc_path.exists():
        print(f"ERROR: Document not found: {doc_path}")
        print("Run with --new to generate from scratch, or --bootstrap to add anchors.")
        sys.exit(1)

    print(f"Opening {doc_path}...")
    doc = Document(str(doc_path))

    # Check if anchors exist
    has_anchors = False
    for para in doc.paragraphs:
        if "[[SECTION:" in para.text:
            has_anchors = True
            break

    if not has_anchors:
        print("No anchors found. Running bootstrap first...")
        doc.save(str(doc_path))
        bootstrap_anchors(doc_path)
        doc = Document(str(doc_path))

    # Parse docker-compose
    compose_services = parse_compose_services()

    # Scan services if requested
    scan_results = None
    if do_scan:
        print("Scanning live services...")
        scan_results = scan_service_health(compose_services)

    # Extract existing changelog entries for preservation
    existing_changelog = []
    cl_start, cl_end = find_anchor_range(doc, "Architecture Changelog")
    if cl_start is not None and cl_end is not None:
        for i in range(cl_start + 1, cl_end):
            text = doc.paragraphs[i].text.strip()
            if text and not text.startswith("Architecture"):
                existing_changelog.append(text)

    # Update each section
    sections_map = {
        "Service Inventory": (generate_service_inventory,),
        "Dependency Matrix": (generate_dependency_matrix,),
        "Ports and Routing": (generate_ports_routing, compose_services),
        "Event Taxonomy": (generate_event_taxonomy,),
        "Enterprise Gate": (generate_enterprise_gate, scan_results),
        "Known Gaps": (generate_known_gaps,),
        "Architecture Changelog": (generate_changelog, existing_changelog),
        "Current Status": (generate_current_status,),
        "Next Actions": (generate_next_actions,),
        "Improvements Required (62% → 100%)": (generate_improvements,),
        "Runtime Flow Diagrams": (generate_runtime_flow,),
    }

    updated = 0
    for section_name, gen_args in sections_map.items():
        fn = gen_args[0]
        args = gen_args[1:] if len(gen_args) > 1 else ()
        print(f"  Updating: {section_name}...")
        if update_section(doc, section_name, fn, *args):
            updated += 1

    # Save
    save_path = out_path or doc_path
    doc.save(str(save_path))
    print(f"\n✅ Updated {updated}/{len(sections_map)} sections")
    print(f"Saved to: {save_path}")
    print(f"Document size: {save_path.stat().st_size / 1024:.0f} KB")

    # Cleanup temp images
    for f in TEMP_DIR.glob("*.png"):
        f.unlink()
    try:
        TEMP_DIR.rmdir()
    except OSError:
        pass


def generate_fresh(out_path: Path):
    """Generate a completely fresh document with all content + anchors."""
    print("Generating fresh document...")

    # First, run the v1 generator to create the base document
    v1_script = BASE_DIR / "generate_aurexis_doc_v1.py"
    if v1_script.exists():
        print("  Running v1 generator for base content...")
        import importlib.util
        spec = importlib.util.spec_from_file_location("v1", str(v1_script))
        v1 = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(v1)
        v1.generate_document()
        print("  v1 base document generated.")
    else:
        print(f"  ⚠ v1 generator not found at {v1_script}")
        print("  Creating minimal document...")
        doc = Document()
        for section in doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
        add_heading_styled(doc, "AUREXIS Architecture Document", 1)
        add_body(doc, "Generated by update_doc.py")
        doc.save(str(out_path))

    # Now bootstrap anchors
    bootstrap_anchors(out_path)

    # Now update all sections
    update_document(out_path)


# ═══════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="AUREXIS Architecture Document Updater v2.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python update_doc.py                    # Update existing doc in-place
  python update_doc.py --bootstrap        # Insert anchors into existing doc
  python update_doc.py --new              # Generate completely fresh doc
  python update_doc.py --out custom.docx  # Write to different file
  python update_doc.py --scan             # Include live health checks
        """)
    parser.add_argument("--bootstrap", action="store_true",
                        help="Insert anchors into existing document")
    parser.add_argument("--new", action="store_true",
                        help="Generate fresh document from scratch")
    parser.add_argument("--out", type=str, default=None,
                        help="Output file path (default: in-place)")
    parser.add_argument("--scan", action="store_true",
                        help="Scan live services for health/ready status")

    args = parser.parse_args()

    out_path = Path(args.out) if args.out else DOC_PATH

    print("=" * 60)
    print(f"AUREXIS Document Updater v{VERSION}")
    print("=" * 60)

    if args.new:
        generate_fresh(out_path)
    elif args.bootstrap:
        if not DOC_PATH.exists():
            print(f"ERROR: {DOC_PATH} not found. Use --new to create from scratch.")
            sys.exit(1)
        bootstrap_anchors(DOC_PATH)
    else:
        update_document(DOC_PATH, out_path, do_scan=args.scan)

    print("=" * 60)


if __name__ == "__main__":
    main()
